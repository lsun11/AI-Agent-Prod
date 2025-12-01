# src/saving/docx_builder.py
import re


def build_docx_document(path: str, query: str, result_text: str) -> None:
    """
    Create a 'paper-style' .docx document:
    - Title = query
    - Section headings for results + recommendations
    """
    from docx import Document

    doc = Document()

    title = query.strip() or "Research Result"
    doc.add_heading(title, level=1)
    doc.add_paragraph("")  # spacer

    lines = result_text.splitlines()

    for line in lines:
        raw = line.rstrip("\n")

        if not raw.strip():
            doc.add_paragraph("")
            continue

        if raw.startswith("📊 Results for:"):
            p = doc.add_paragraph(raw)
            p.style = "Intense Quote" if "Intense Quote" in doc.styles else "Normal"
            continue

        if raw.strip().startswith("**Recommendations / Analysis:**") or \
           raw.strip().startswith("Recommendations / Analysis:"):
            clean = raw.replace("**", "").strip()
            doc.add_heading(clean, level=2)
            continue

        if raw.strip().startswith("## "):
            clean = raw.strip().lstrip("#").strip()
            doc.add_heading(clean, level=2)
            continue

        if raw.strip().startswith("### "):
            clean = raw.strip().lstrip("#").strip()
            doc.add_heading(clean, level=3)
            continue

        if raw.lstrip().startswith("- "):
            doc.add_paragraph(raw.lstrip()[2:], style="List Bullet")
            continue

        if re.match(r"^\d+\.\s", raw):
            doc.add_paragraph(raw, style="List Number")
            continue

        doc.add_paragraph(raw)

    doc.save(path)
